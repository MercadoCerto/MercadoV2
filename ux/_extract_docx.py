import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

d = Document(r'C:\Users\david.6961\mercadocerto\MercadoCerto-main\startup\ux\MercadoCerto_Trabalho2_UX.docx')

for p in d.paragraphs:
    print(p.text)

print("\n\n=== TABELAS ===\n")
for i, tab in enumerate(d.tables):
    print(f"\n--- Tabela {i+1} ---")
    for row in tab.rows:
        cells = [c.text.strip() for c in row.cells]
        print(" | ".join(cells))

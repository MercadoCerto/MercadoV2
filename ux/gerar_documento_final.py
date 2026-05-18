# -*- coding: utf-8 -*-
"""
Gera a versão refinada da documentação UX do MercadoCerto (Trabalho 2).
Saída: MercadoCerto_Trabalho2_UX_Final.docx
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

PATH_OUT = r"C:\Users\david.6961\mercadocerto\MercadoCerto-main\startup\ux\MercadoCerto_Trabalho2_UX_Final.docx"

doc = Document()

# ---------- Estilos básicos ----------
estilo_normal = doc.styles['Normal']
estilo_normal.font.name = 'Times New Roman'
estilo_normal.font.size = Pt(12)

for nivel in [1, 2, 3]:
    h = doc.styles[f'Heading {nivel}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h.font.bold = True
    h.font.size = Pt(14 if nivel == 1 else 12)

# Margens ABNT
secao = doc.sections[0]
secao.top_margin = Cm(3)
secao.bottom_margin = Cm(2)
secao.left_margin = Cm(3)
secao.right_margin = Cm(2)


def add_par(texto, *, bold=False, align=None, espaco_antes=0, espaco_depois=6,
            primeira_linha=None, italico=False, justificar=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    elif justificar:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(espaco_antes)
    pf.space_after = Pt(espaco_depois)
    if primeira_linha is not None:
        pf.first_line_indent = Cm(primeira_linha)
    run = p.add_run(texto)
    run.bold = bold
    run.italic = italico
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_heading(texto, nivel):
    h = doc.add_heading(texto, level=nivel)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.space_before = Pt(18 if nivel == 1 else 12)
    h.paragraph_format.space_after = Pt(8)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.size = Pt(14 if nivel == 1 else 12)
    return h


def add_bullet(texto):
    p = doc.add_paragraph(texto, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p


def add_pagebreak():
    doc.add_page_break()


def add_capa_par(texto, bold=False, size=12, alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
                 espaco_antes=0, espaco_depois=6):
    p = doc.add_paragraph()
    p.alignment = alinhamento
    p.paragraph_format.space_before = Pt(espaco_antes)
    p.paragraph_format.space_after = Pt(espaco_depois)
    run = p.add_run(texto)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p


# ==========================================================
# CAPA
# ==========================================================
add_capa_par("CENTRO UNIVERSITÁRIO SENAC", bold=True, size=12, espaco_antes=0)
add_capa_par("ANÁLISE E DESENVOLVIMENTO DE SISTEMAS", bold=True, size=12)
add_capa_par("DISCIPLINA: UX – ARQUITETURA E USABILIDADE", bold=True, size=12, espaco_depois=24)

for _ in range(2):
    add_capa_par(" ")

add_capa_par("DAVID SOUZA", bold=True, size=12)
add_capa_par("LUCAS MICHELS", bold=True, size=12)
add_capa_par("MATHEUS BOLLER", bold=True, size=12)
add_capa_par("PEDRO KARATCHUK", bold=True, size=12, espaco_depois=36)

for _ in range(3):
    add_capa_par(" ")

add_capa_par("MERCADOCERTO", bold=True, size=16)
add_capa_par("Projeto de Interface Web Centrado no Usuário:", bold=True, size=14)
add_capa_par("Melhorias de UX e Aplicação dos Conceitos de IHC", bold=True, size=14, espaco_depois=36)

for _ in range(6):
    add_capa_par(" ")

add_capa_par("Curitiba", bold=True, size=12)
add_capa_par("2026", bold=True, size=12)
add_pagebreak()

# ==========================================================
# FOLHA DE ROSTO
# ==========================================================
add_capa_par("DAVID SOUZA", bold=True, size=12)
add_capa_par("LUCAS MICHELS", bold=True, size=12)
add_capa_par("MATHEUS BOLLER", bold=True, size=12)
add_capa_par("PEDRO KARATCHUK", bold=True, size=12, espaco_depois=48)

for _ in range(4):
    add_capa_par(" ")

add_capa_par("MERCADOCERTO", bold=True, size=16)
add_capa_par("Projeto de Interface Web Centrado no Usuário:", bold=True, size=13)
add_capa_par("Melhorias de UX e Aplicação dos Conceitos de IHC", bold=True, size=13, espaco_depois=36)

p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_desc.paragraph_format.left_indent = Cm(8)
p_desc.paragraph_format.line_spacing = 1.0
p_desc.paragraph_format.space_before = Pt(24)
p_desc.paragraph_format.space_after = Pt(24)
run_desc = p_desc.add_run(
    "Trabalho 2 apresentado à disciplina de UX – Arquitetura e Usabilidade, "
    "do curso de Análise e Desenvolvimento de Sistemas, como parte dos "
    "requisitos para avaliação parcial.\n\n"
    "Professor: Carlos Alberto Pedroso Junior."
)
run_desc.font.name = 'Times New Roman'
run_desc.font.size = Pt(11)

for _ in range(6):
    add_capa_par(" ")

add_capa_par("Curitiba", bold=True, size=12)
add_capa_par("2026", bold=True, size=12)
add_pagebreak()

# ==========================================================
# SUMÁRIO
# ==========================================================
add_heading("SUMÁRIO", 1)
sumario = [
    ("1 INTRODUÇÃO", "4"),
    ("2 FUNDAMENTAÇÃO TEÓRICA", "5"),
    ("2.1 Interação Humano-Computador", "5"),
    ("2.2 Qualidade de Uso", "6"),
    ("2.3 Interface", "6"),
    ("2.4 Affordance", "7"),
    ("3 METODOLOGIA", "8"),
    ("4 PARTE 1 – DOCUMENTAÇÃO DAS MELHORIAS", "9"),
    ("4.1 Tipografia e legibilidade", "9"),
    ("4.2 Contraste e cores", "10"),
    ("4.3 Navegação e botão Voltar", "10"),
    ("4.4 Tela do mapa interativo", "11"),
    ("4.5 Separação de perfis: Consumidor e Comerciante", "11"),
    ("4.6 Validação de preços (anti-fraude)", "12"),
    ("4.7 Cadastro de produto – UX do formulário", "13"),
    ("4.8 Lista de compras", "13"),
    ("4.9 Tela inicial (Home)", "14"),
    ("4.10 Responsividade", "14"),
    ("4.11 Avaliações e ranking", "15"),
    ("4.12 Correções de bugs", "15"),
    ("5 PARTE 2 – APLICAÇÃO DOS CONCEITOS-CHAVE DE IHC", "16"),
    ("5.1 Interação", "16"),
    ("5.2 Affordance", "18"),
    ("5.3 Interface", "19"),
    ("5.4 Qualidade de Uso", "21"),
    ("6 QUADRO RESUMO – MELHORIAS E CONCEITOS APLICADOS", "23"),
    ("7 ANÁLISE CRÍTICA E CONSIDERAÇÕES FINAIS", "24"),
    ("REFERÊNCIAS", "26"),
]
for item, pag in sumario:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), alignment=2)
    run = p.add_run(item + "\t" + pag)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
add_pagebreak()

# ==========================================================
# 1 INTRODUÇÃO
# ==========================================================
add_heading("1 INTRODUÇÃO", 1)

add_par(
    "Este trabalho documenta a segunda fase do projeto MercadoCerto, plataforma web "
    "de comparação de preços e avaliação de mercados desenvolvida ao longo da disciplina "
    "de UX – Arquitetura e Usabilidade. A proposta desta etapa é revisar a interface "
    "construída na primeira entrega, aplicando de forma deliberada os conceitos de "
    "interação, affordance, interface e qualidade de uso estudados em aula.",
    primeira_linha=1.25
)

add_par(
    "A motivação do MercadoCerto surge da percepção de um problema cotidiano: o consumidor "
    "frequentemente não sabe em qual mercado da região um item está com o preço mais "
    "vantajoso, e acaba percorrendo lojas distintas sem uma referência clara. A plataforma "
    "busca atender essa lacuna ao reunir, em um único ambiente, comparação de preços, "
    "avaliações de estabelecimentos e listas de compras inteligentes, organizadas a partir "
    "da localização do usuário.",
    primeira_linha=1.25
)

add_par(
    "Na primeira entrega, o foco esteve no desenvolvimento funcional. A interface, ainda "
    "que coerente em termos de estrutura, apresentava limitações que ficaram evidentes na "
    "avaliação recebida. Problemas de tipografia, contraste, navegação, validação de dados "
    "e diferenciação entre perfis de usuário comprometiam a clareza da experiência. A partir "
    "desse diagnóstico, esta segunda entrega assume um caráter de refinamento: cada "
    "apontamento foi analisado e convertido em uma ação concreta de melhoria, com base em "
    "fundamentos teóricos da Interação Humano-Computador (IHC).",
    primeira_linha=1.25
)

add_par(
    "O documento está organizado em sete seções principais. Após esta introdução, a Seção 2 "
    "apresenta a fundamentação teórica, retomando os conceitos centrais discutidos em aula. "
    "A Seção 3 descreve a metodologia adotada para coletar e tratar o feedback. A Seção 4 "
    "(Parte 1 do trabalho) documenta as melhorias implementadas, organizadas por categoria. "
    "A Seção 5 (Parte 2 do trabalho) faz a ligação explícita entre cada melhoria e os "
    "conceitos-chave de IHC. A Seção 6 traz um quadro-resumo dessa correspondência, e a "
    "Seção 7 apresenta uma análise crítica do processo e dos resultados alcançados.",
    primeira_linha=1.25
)

add_pagebreak()

# ==========================================================
# 2 FUNDAMENTAÇÃO TEÓRICA
# ==========================================================
add_heading("2 FUNDAMENTAÇÃO TEÓRICA", 1)
add_par(
    "Esta seção retoma os conceitos centrais da disciplina que sustentam as decisões de "
    "projeto descritas mais adiante. Em vez de tratar IHC como um rótulo isolado, buscou-se "
    "compreender como cada conceito se articula com a prática do desenvolvimento de "
    "interfaces web.",
    primeira_linha=1.25
)

add_heading("2.1 Interação Humano-Computador", 2)
add_par(
    "A Interação Humano-Computador, segundo Winograd (2003), é o estudo do processo de "
    "comunicação entre humanos e sistemas computacionais. Trata-se de uma área "
    "interdisciplinar, que reúne Computação, Psicologia Cognitiva e Ergonomia, com o "
    "objetivo de tornar os sistemas mais úteis, eficientes e agradáveis aos usuários.",
    primeira_linha=1.25
)
add_par(
    "Mayhew (1992) descreve a interação como um processo dinâmico organizado em três fases: "
    "ler-examinar, pensar e responder. O usuário observa o estado do sistema, interpreta a "
    "informação recebida e decide qual ação executar; em seguida, o sistema processa essa "
    "entrada e devolve uma resposta, reiniciando o ciclo. A interface, nesse modelo, atua "
    "como o ponto de contato em que o controle alterna entre humano e máquina.",
    primeira_linha=1.25
)
add_par(
    "Norman (2013), por sua vez, descreve a ação do usuário como um ciclo composto por "
    "etapas de percepção, decisão, execução e avaliação. Para que a experiência seja "
    "satisfatória, esse ciclo precisa fluir sem rupturas, com feedback claro a cada etapa. "
    "Esse referencial guiou diretamente as escolhas relacionadas ao fluxo de tarefas do "
    "MercadoCerto.",
    primeira_linha=1.25
)

add_heading("2.2 Qualidade de Uso", 2)
add_par(
    "A norma ISO 9241-11 define qualidade de uso como a medida em que um sistema permite "
    "que usuários atinjam objetivos específicos com eficácia, eficiência e satisfação em um "
    "contexto de uso determinado. Nielsen (1993) amplia essa visão ao propor cinco atributos "
    "da usabilidade: facilidade de aprendizado, eficiência, memorabilidade, baixa taxa de "
    "erros e satisfação subjetiva.",
    primeira_linha=1.25
)
add_par(
    "O material da disciplina acrescenta a esse conjunto a dimensão de acessibilidade, "
    "compreendida como a capacidade do sistema de ser utilizado por pessoas com diferentes "
    "habilidades físicas, sensoriais ou cognitivas. Esses quatro critérios — eficácia, "
    "eficiência, satisfação e acessibilidade — foram adotados como referência central para "
    "a avaliação das melhorias implementadas.",
    primeira_linha=1.25
)
add_par(
    "Importa observar a distinção trazida em aula entre um sistema tecnicamente robusto e um "
    "sistema com alta qualidade de uso. O primeiro pode ser estável, performático e bem "
    "documentado, mas ainda assim oferecer uma experiência frustrante ao usuário final. A "
    "qualidade de uso desloca o foco do funcionamento interno para a experiência vivida.",
    primeira_linha=1.25
)

add_heading("2.3 Interface", 2)
add_par(
    "A interface é o canal de mediação entre usuário e sistema. Ela engloba os elementos "
    "visuais, sonoros e táteis por meio dos quais a interação ocorre. Conforme discutido em "
    "sala, há vários tipos de interface — gráficas, textuais, conversacionais e físicas — e "
    "a escolha de cada uma está condicionada ao contexto de uso, ao dispositivo e ao perfil "
    "do usuário.",
    primeira_linha=1.25
)
add_par(
    "No caso do MercadoCerto, a interface adotada é gráfica (GUI), com versões para desktop "
    "e dispositivos móveis. Três atributos foram priorizados no refinamento: consistência "
    "visual e semântica, organização lógica com hierarquia clara e uso de padrões de design "
    "já reconhecidos pelos usuários. Esses atributos dialogam diretamente com as Heurísticas "
    "de Nielsen (1994), especialmente as referentes à visibilidade do estado do sistema, "
    "consistência, prevenção de erros e reconhecimento em vez de memorização.",
    primeira_linha=1.25
)

add_heading("2.4 Affordance", 2)
add_par(
    "O conceito de affordance foi originalmente proposto por Gibson (1979), no campo da "
    "psicologia ecológica, para descrever as possibilidades de ação que o ambiente oferece "
    "a um agente. Norman (1988, 2013) trouxe o termo para o design, distinguindo as "
    "affordances reais (o que o objeto de fato permite) das affordances percebidas (o que "
    "o usuário imagina que pode fazer com ele).",
    primeira_linha=1.25
)
add_par(
    "No contexto de interfaces digitais, costuma-se classificar as affordances em três "
    "tipos: perceptível (a aparência comunica claramente a função, como um botão em relevo); "
    "falsa affordance (o elemento sugere uma ação que, na prática, não existe, como um texto "
    "sublinhado que não leva a lugar nenhum); e oculta (a função existe, mas nenhuma pista "
    "visual a indica). O trabalho de refinamento do MercadoCerto procurou reforçar as "
    "affordances perceptíveis e eliminar as falsas affordances identificadas no review.",
    primeira_linha=1.25
)

add_pagebreak()

# ==========================================================
# 3 METODOLOGIA
# ==========================================================
add_heading("3 METODOLOGIA", 1)
add_par(
    "O processo de refinamento foi conduzido em quatro etapas. A primeira consistiu na "
    "coleta de feedback sobre a versão entregue no Trabalho 1, a partir de duas fontes: "
    "(a) os comentários formais do professor da disciplina e (b) uma sessão de revisão "
    "entre os integrantes do grupo, simulando uma avaliação de pares. Em ambos os casos, "
    "os apontamentos foram registrados de forma estruturada, separando aspectos visuais, "
    "de navegação, de lógica de negócio e de acessibilidade.",
    primeira_linha=1.25
)
add_par(
    "Na segunda etapa, cada apontamento foi classificado segundo sua natureza e prioridade. "
    "Foram considerados críticos os problemas que comprometiam a usabilidade básica (como "
    "contraste insuficiente e ausência de validação de dados) e secundários os ajustes "
    "voltados principalmente à estética e ao polimento.",
    primeira_linha=1.25
)
add_par(
    "A terceira etapa foi a definição das ações corretivas, sempre com fundamento em pelo "
    "menos um dos quatro conceitos-chave da disciplina. Para cada problema, foi feito o "
    "exercício de identificar a qual princípio de IHC o ajuste se vinculava, evitando "
    "decisões meramente estéticas.",
    primeira_linha=1.25
)
add_par(
    "A quarta etapa, finalmente, envolveu a implementação das mudanças no código da "
    "aplicação e a documentação reflexiva apresentada neste relatório. As decisões foram "
    "registradas também em comentários dos arquivos CSS e nos templates HTML, de modo a "
    "preservar a rastreabilidade entre a justificativa conceitual e o código entregue.",
    primeira_linha=1.25
)

add_pagebreak()

# ==========================================================
# 4 PARTE 1 — DOCUMENTAÇÃO DAS MELHORIAS
# ==========================================================
add_heading("4 PARTE 1 – DOCUMENTAÇÃO DAS MELHORIAS", 1)
add_par(
    "Esta seção apresenta as melhorias implementadas a partir do feedback recebido, "
    "organizadas em doze categorias. Para cada uma, descreve-se o problema original, as "
    "ações adotadas e o efeito esperado na experiência do usuário.",
    primeira_linha=1.25
)

# 4.1
add_heading("4.1 Tipografia e legibilidade", 2)
add_par(
    "A tipografia foi um dos pontos mais recorrentes nos comentários da avaliação. A queixa "
    "central indicava que a fonte utilizada nos títulos e na interface possuía caráter "
    "excessivamente decorativo, o que dificultava a leitura, sobretudo para o público-alvo "
    "da plataforma: consumidores em geral, com destaque para pessoas mais velhas e usuários "
    "com baixa visão.",
    primeira_linha=1.25
)
add_par("Ações implementadas:", bold=True, primeira_linha=1.25)
add_bullet("Substituição da fonte decorativa pela família Poppins (Google Fonts), uma sans-serif geométrica de boa legibilidade em tela; a fonte anterior permaneceu apenas no logotipo, preservando a identidade visual da marca;")
add_bullet("Revisão da escala tipográfica: títulos principais em 20px, subtítulos em 16px e corpo de texto em 15px, eliminando o excesso de tamanho apontado;")
add_bullet("Padronização dos pesos tipográficos — Bold (700) para títulos, SemiBold (600) para subtítulos, Regular (400) para corpo e Medium (500) para labels e botões —, criando uma hierarquia visual previsível.")

# 4.2
add_heading("4.2 Contraste e cores", 2)
add_par(
    "O segundo ponto crítico identificado foi o contraste. Textos em cinza claro sobre fundo "
    "claro e combinações de roxo sobre roxo comprometiam tanto a leitura quanto a "
    "acessibilidade. Embora a estética original buscasse uma identidade marcante, o "
    "resultado tornava trechos inteiros da interface praticamente ilegíveis.",
    primeira_linha=1.25
)
add_par("Ações implementadas:", bold=True, primeira_linha=1.25)
add_bullet("Padronização do fundo em tons claros — branco (#FFFFFF) para cards e lavanda suave (#F5F3FF) para o fundo geral —, mais coerente com a proposta de um sistema de compras;")
add_bullet("Adoção do tom #1A1A2E para o texto principal sobre fundos claros, atingindo razão de contraste superior a 7:1, em conformidade com o nível AAA da WCAG 2.1;")
add_bullet("Reformulação das combinações roxo-sobre-roxo: nos pontos em que o texto era roxo claro sobre fundo roxo, passou-se a usar texto escuro sobre fundo claro ou texto branco sobre fundo roxo sólido (#7C3AED), sempre respeitando o contraste mínimo;")
add_bullet("Ampliação do espaço entre o texto “Cadastre-se gratuitamente” e os ícones de criação de conta, com padding de 16px e margin-top de 12px, eliminando o efeito de elementos sobrepostos.")

# 4.3
add_heading("4.3 Navegação e botão Voltar", 2)
add_par(
    "Identificou-se a ausência de botão de voltar em diversas telas internas, o que obrigava "
    "o usuário a recorrer ao botão do navegador. Em telas onde o botão existia, o "
    "redirecionamento estava incorreto, conduzindo a páginas que não correspondiam ao passo "
    "anterior do fluxo.",
    primeira_linha=1.25
)
add_par("Ações implementadas:", bold=True, primeira_linha=1.25)
add_bullet("Padronização do botão “Voltar” em todas as telas internas, posicionado no canto superior esquerdo, logo abaixo do header, com ícone de seta e rótulo textual, no estilo de breadcrumb;")
add_bullet("Correção dos redirecionamentos existentes, de modo que cada botão conduza à tela hierarquicamente anterior no fluxo;")
add_bullet("Implementação de uma lógica de navegação com rotas explícitas como caminho principal e history.back() como fallback, tornando o comportamento previsível mesmo em navegação com múltiplas abas.")

# 4.4
add_heading("4.4 Tela do mapa interativo", 2)
add_par(
    "O mapa abria com a visualização do Brasil inteiro, sem qualquer filtro inicial, o que "
    "pouco auxiliava o usuário no primeiro contato. Além disso, ao clicar em um mercado, "
    "não havia caminho direto para acessar a página detalhada do estabelecimento.",
    primeira_linha=1.25
)
add_par("Ações implementadas:", bold=True, primeira_linha=1.25)
add_bullet("Centralização automática do mapa na posição do usuário, com base na API de Geolocation do navegador, mediante consentimento;")
add_bullet("Definição de um raio padrão de 50 km, evitando a exibição inicial em escala nacional;")
add_bullet("Inclusão de um campo de busca manual por região (cidade, bairro ou CEP), atendendo aos usuários que desejem consultar mercados em outras localidades;")
add_bullet("Adição, no popup do marcador, de um botão “Ver detalhes”, que conduz diretamente à página do mercado.")

# 4.5
add_heading("4.5 Separação de perfis: Consumidor e Comerciante", 2)
add_par(
    "Um dos problemas mais críticos identificados foi a ausência de diferenciação entre "
    "perfis de usuário. Na versão anterior, qualquer pessoa cadastrada podia inserir "
    "produtos em qualquer mercado, o que abria espaço para informações fraudulentas — como, "
    "por exemplo, o cadastro deliberado de preços absurdos em concorrentes. Sem essa "
    "separação, a confiabilidade dos dados ficava comprometida.",
    primeira_linha=1.25
)
add_par("Ações implementadas:", bold=True, primeira_linha=1.25)
add_bullet("Criação de dois tipos de cadastro distintos: Consumidor (com CPF) e Comerciante (com CNPJ), com validação do documento no momento do registro;")
add_bullet("Restrição das funções de cadastro e edição de produtos exclusivamente ao comerciante vinculado ao mercado, impedindo que consumidores comuns publiquem preços;")
add_bullet("Disponibilização de um Painel do Comércio, com menu próprio: Meus Mercados, Meus Produtos, Publicar Preço, Avaliações Recebidas e Dashboard de Reputação;")
add_bullet("Suporte a redes com várias unidades: no caso de franquias, o comerciante seleciona em qual filial o produto será cadastrado, evitando confusão entre lojas da mesma rede;")
add_bullet("Remoção do botão “Cadastrar Produto” da listagem pública de produtos, mantendo essa ação apenas no painel restrito ao comerciante.")

# 4.6
add_heading("4.6 Validação de preços (anti-fraude)", 2)
add_par(
    "Mesmo após a separação de perfis, restava o risco de comerciantes registrarem preços "
    "incorretos por engano ou de forma deliberada. Para mitigar essa possibilidade, foram "
    "implementados mecanismos de validação no momento da publicação do preço.",
    primeira_linha=1.25
)
add_par("Ações implementadas:", bold=True, primeira_linha=1.25)
add_bullet("Obrigatoriedade do upload da foto da etiqueta de preço no cadastro, funcionando como comprovação visual do valor informado;")
add_bullet("Integração de um leitor de código de barras via câmera do dispositivo, com o uso da biblioteca QuaggaJS, que vincula o produto ao código do fabricante de forma automática;")
add_bullet("Inclusão de uma validação de faixa de preço: quando o valor informado destoa significativamente da média histórica do produto (variação acima de 50%), o sistema exibe um alerta e solicita confirmação do comerciante;")
add_bullet("Previsão, em uma próxima iteração, de leitura automática de preço e código de barras a partir da foto da etiqueta, com técnicas de OCR e visão computacional.")

# 4.7
add_heading("4.7 Cadastro de produto – UX do formulário", 2)
add_par(
    "O formulário de cadastro de produto apresentava vários problemas de usabilidade. O "
    "campo de preço incrementava apenas um centavo por clique, o que tornava inviável "
    "definir valores altos. A validação misturava vírgula e ponto, gerando erros de "
    "processamento, e não havia formatação automática do valor digitado.",
    primeira_linha=1.25
)
add_par("Ações implementadas:", bold=True, primeira_linha=1.25)
add_bullet("Substituição do input numérico padrão por um campo monetário com máscara automática: ao digitar “1599”, o sistema exibe “R$ 15,99”, eliminando ambiguidades de formato;")
add_bullet("Adição, ao lado do campo, de botões de incremento rápido (+R$ 0,10, +R$ 1,00 e +R$ 10,00), agilizando ajustes em valores maiores;")
add_bullet("Aceitação tanto de vírgula quanto de ponto como separadores decimais, com normalização interna para o formato correto antes do envio ao backend.")

# 4.8
add_heading("4.8 Lista de compras", 2)
add_par(
    "A funcionalidade de lista de compras possuía duas limitações principais. A primeira "
    "era a impossibilidade de excluir um item específico — só era permitido remover a lista "
    "inteira. A segunda era a ausência de modo de edição após a criação da lista, obrigando "
    "o usuário a recomeçar o processo em caso de erro.",
    primeira_linha=1.25
)
add_par("Ações implementadas:", bold=True, primeira_linha=1.25)
add_bullet("Inclusão de um botão individual de exclusão em cada item (ícone de lixeira), com tooltip de confirmação para evitar remoções acidentais;")
add_bullet("Implementação de um modo de edição que permite alterar quantidades, renomear itens e reordená-los após a criação da lista;")
add_bullet("Correção do fluxo de cadastro item-por-linha: vários produtos digitados em linhas separadas passam a ser agrupados corretamente em uma única lista.")

# 4.9
add_heading("4.9 Tela inicial (Home)", 2)
add_par(
    "A tela inicial foi descrita pelos avaliadores como visualmente esvaziada e sem "
    "hierarquia. As opções apareciam dispostas sem priorização clara, tanto para usuários "
    "logados quanto para visitantes, o que dificultava o primeiro contato com a plataforma.",
    primeira_linha=1.25
)
add_par("Ações implementadas:", bold=True, primeira_linha=1.25)
add_bullet("Redesenho da home com um carrossel de destaques no topo (promoções, mercados bem avaliados e produtos populares), funcionando como ponto focal da página;")
add_bullet("Reorganização do grid de ações rápidas em dois níveis de hierarquia: ações principais (Comparar Preços, Minha Lista, Mercados Próximos) em cards de maior dimensão e ações secundárias (Ranking, Mapa, Cadastrar) em cards menores logo abaixo;")
add_bullet("Personalização da home para usuários autenticados, com saudação, resumo das listas recentes e sugestões baseadas no histórico, em vez da exibição indiferenciada de todos os menus.")

# 4.10
add_heading("4.10 Responsividade", 2)
add_par(
    "O layout em dispositivos móveis já funcionava parcialmente, mas algumas seções não se "
    "adaptavam corretamente, quebrando em telas pequenas. Como parte significativa do "
    "público-alvo realiza compras consultando o celular dentro do próprio mercado, esse "
    "ponto foi tratado como prioridade.",
    primeira_linha=1.25
)
add_par("Ações implementadas:", bold=True, primeira_linha=1.25)
add_bullet("Revisão dos breakpoints CSS para smartphones (320px–480px), tablets (481px–768px) e desktops (acima de 769px);")
add_bullet("Conversão das tabelas de comparação de preços em cards empilhados nas telas pequenas, evitando rolagem horizontal;")
add_bullet("Ajuste do menu hamburger para funcionar de forma consistente em todas as páginas, com animação suave de abertura e fechamento.")

# 4.11
add_heading("4.11 Avaliações e ranking", 2)
add_par(
    "A fórmula original do ranking dava o mesmo peso a um mercado com uma única avaliação "
    "de cinco estrelas e a outro com centenas de avaliações positivas, o que distorcia o "
    "resultado e prejudicava a confiabilidade da lista.",
    primeira_linha=1.25
)
add_par("Ações implementadas:", bold=True, primeira_linha=1.25)
add_bullet("Adoção de média ponderada bayesiana, que combina a nota média e o número de avaliações, evitando que estabelecimentos com poucas avaliações ocupem o topo da lista;")
add_bullet("Exibição do número total de avaliações ao lado da nota média em todos os contextos (lista de mercados, ranking e cards), conferindo maior transparência ao usuário;")
add_bullet("Definição de critério de desempate por volume de avaliações, em substituição à ordenação alfabética que ocorria por padrão.")

# 4.12
add_heading("4.12 Correções de bugs", 2)
add_par(
    "Além das melhorias estruturais, vários bugs pontuais foram identificados no review e "
    "corrigidos nesta fase:",
    primeira_linha=1.25
)
add_bullet("Correção dos botões de voltar que redirecionavam para a tela incorreta;")
add_bullet("Inclusão de fallback para imagens antigas armazenadas no banco: quando a URL original falha, o sistema exibe automaticamente um placeholder;")
add_bullet("Reforço da validação no cadastro de conta, com checagens no frontend (HTML5 required + JavaScript) e no backend (Bean Validation do Spring), recusando campos nulos e e-mails malformados.")

add_pagebreak()

# ==========================================================
# 5 PARTE 2 — APLICAÇÃO DOS CONCEITOS-CHAVE
# ==========================================================
add_heading("5 PARTE 2 – APLICAÇÃO DOS CONCEITOS-CHAVE DE IHC", 1)
add_par(
    "Nesta seção, as melhorias descritas anteriormente são relidas a partir dos quatro "
    "conceitos-chave previstos no enunciado do trabalho: interação, affordance, interface e "
    "qualidade de uso. O objetivo não é repetir os ajustes, mas mostrar de que forma cada "
    "decisão se ancora em princípios teóricos e em referências da literatura de IHC.",
    primeira_linha=1.25
)

# 5.1 INTERAÇÃO
add_heading("5.1 Interação", 2)
add_heading("5.1.1 Definição do fluxo de tarefas do usuário", 3)
add_par(
    "A separação de perfis (Seção 4.5) reestruturou os fluxos de tarefas da plataforma. Na "
    "versão anterior, havia um único fluxo genérico, em que todos os usuários podiam "
    "executar qualquer ação — configuração incompatível com o modelo de negócio. Após a "
    "revisão, dois fluxos distintos passaram a coexistir:",
    primeira_linha=1.25
)
add_par(
    "Fluxo do Consumidor: Home → Buscar Produtos → Comparar Preços → Adicionar à Lista → "
    "Otimizar Lista → Plano de Compras. Cada etapa é sequencial e oferece feedback visual "
    "nas transições (loading spinner, toast de confirmação).",
    primeira_linha=1.25
)
add_par(
    "Fluxo do Comerciante: Login (CNPJ) → Painel do Comércio → Selecionar Mercado → "
    "Cadastrar Produto com foto da etiqueta → Publicar Preço → Acompanhar Avaliações. Aqui "
    "o fluxo é guiado: cada etapa indica explicitamente qual é a próxima ação prevista.",
    primeira_linha=1.25
)
add_par(
    "Essa reestruturação dialoga com os Sete Estágios da Ação de Norman (2013), nos quais o "
    "usuário percorre etapas de formação do objetivo, planejamento, execução e avaliação. "
    "Ao tornar os fluxos explícitos, a interface reduz o esforço cognitivo associado a cada "
    "uma dessas etapas.",
    primeira_linha=1.25
)

add_heading("5.1.2 Ciclos de ação e feedback", 3)
add_par(
    "Toda ação do usuário passa a gerar uma resposta visível e imediata do sistema. Esse "
    "ciclo de ação-feedback aparece em vários pontos:",
    primeira_linha=1.25
)
add_bullet("Cadastro de produto com foto: ao subir a etiqueta, o sistema apresenta preview da imagem, barra de progresso e confirmação visual (toast verde “Produto cadastrado com sucesso”), completando o ciclo de ação;")
add_bullet("Validação de preço fora da faixa: quando o comerciante informa um valor 50% acima da média histórica, o sistema exibe alerta amarelo com a referência e pede confirmação — um feedback preventivo que interrompe o ciclo para evitar erro;")
add_bullet("Busca no mapa: ao digitar uma cidade ou bairro, o autocomplete sugere opções em tempo real; ao selecionar uma sugestão, o mapa anima até a nova localização e os marcadores aparecem progressivamente;")
add_bullet("Exclusão de item da lista: ao acionar a lixeira, surge um tooltip de confirmação “Remover item?” com opções Sim/Não, seguido de animação de remoção e toast informativo, prevenindo ações irreversíveis por engano.")

add_heading("5.1.3 Respostas visíveis do sistema", 3)
add_par(
    "A revisão da navegação (Seção 4.3) ilustra bem o princípio de visibilidade do estado "
    "do sistema. O botão de voltar funciona como uma “saída de emergência” sempre "
    "visível, atendendo à Heurística 3 de Nielsen (1994), e o destaque do link ativo no "
    "header sinaliza ao usuário em que ponto da hierarquia ele se encontra, em consonância "
    "com a Heurística 1 (visibilidade do status do sistema).",
    primeira_linha=1.25
)
add_par(
    "Spinners durante operações assíncronas e toasts contextuais — verde para sucesso, "
    "vermelho para erro, amarelo para aviso — garantem que o sistema nunca permaneça em "
    "silêncio durante uma interação prolongada. Esse princípio, embora simples, é um dos "
    "pilares do design de interação segundo Cooper et al. (2014).",
    primeira_linha=1.25
)

# 5.2 AFFORDANCE
add_heading("5.2 Affordance", 2)
add_heading("5.2.1 Elementos que indicam claramente sua função", 3)
add_par(
    "O conceito de affordance, na acepção adotada por Norman (2013), refere-se às "
    "propriedades percebidas de um objeto que sugerem sua utilização. Diversas melhorias do "
    "MercadoCerto reforçam essas affordances perceptíveis:",
    primeira_linha=1.25
)
add_bullet("O botão “Voltar” combina ícone de seta e rótulo textual. A duplicidade de pistas — direção indicada pelo ícone e confirmação pelo texto — reduz a ambiguidade comum em botões apenas icônicos;")
add_bullet("Os botões de incremento de preço (+R$ 0,10, +R$ 1,00 e +R$ 10,00) explicitam, no próprio rótulo, o valor a ser adicionado. A forma retangular com borda e cor de destaque comunica a affordance de “clicável”;")
add_bullet("O ícone de lixeira para excluir itens aciona uma metáfora visual universalmente reconhecida; a cor vermelha reforça a semântica de ação destrutiva, prevenindo equívocos;")
add_bullet("O campo de busca no mapa adota ícone de lupa e placeholder “Buscar cidade ou bairro…”, aliando ícone e texto para evidenciar a função de pesquisa.")

add_heading("5.2.2 Eliminação de falsas affordances", 3)
add_par(
    "O redesenho da home (Seção 4.9) e a reformulação do cadastro de produtos (Seções 4.5 "
    "e 4.7) eliminaram falsas affordances apontadas no review:",
    primeira_linha=1.25
)
add_bullet("Cards informativos que aparentavam ser clicáveis sem possuir essa função tiveram seu efeito hover removido. O cursor pointer e a elevação de sombra agora ocorrem apenas em elementos efetivamente interativos;")
add_bullet("O botão “Cadastrar Produto” deixou a listagem pública de produtos. Sua presença naquele contexto sugeria, ao consumidor, uma ação que de fato não lhe cabia, configurando uma falsa affordance;")
add_bullet("Textos sublinhados que não constituíam hyperlinks foram convertidos em negrito ou itálico. O sublinhado ficou reservado exclusivamente aos links reais, em respeito à convenção de uso da web.")

add_heading("5.2.3 Cores, formas e metáforas visuais", 3)
add_par(
    "A revisão da paleta (Seção 4.2) utiliza as cores como signifiers, isto é, como "
    "elementos que comunicam significado para além da estética:",
    primeira_linha=1.25
)
add_bullet("Verde (#28A745) para confirmar e sinalizar sucesso: botões Salvar, Cadastrar e toasts de confirmação;")
add_bullet("Vermelho (#E03C3C) para ações destrutivas e alertas: botão Excluir, preços em destaque, mensagens de erro;")
add_bullet("Roxo (#7C3AED) para ações primárias e elementos de navegação: botão de ação principal, links do header e estados ativos;")
add_bullet("Amarelo (#F5A623) para avisos e destaques: alertas de preço fora da faixa, estrelas de avaliação, badges de destaque.")
add_par(
    "Essa codificação cromática se apoia no princípio de mapeamento natural (Norman, 2013): "
    "vermelho e verde carregam significados culturalmente estabelecidos (perigo/parar e "
    "seguro/prosseguir), o que reduz a carga cognitiva e acelera o reconhecimento das "
    "funções.",
    primeira_linha=1.25
)

# 5.3 INTERFACE
add_heading("5.3 Interface", 2)
add_heading("5.3.1 Consistência visual e semântica", 3)
add_par(
    "A padronização tipográfica (Seção 4.1) e a correção do contraste (Seção 4.2) garantem "
    "consistência visual em toda a aplicação. Os elementos seguem um pequeno Design System, "
    "definido a partir de tokens CSS:",
    primeira_linha=1.25
)
add_bullet("A fonte Poppins é adotada em todos os contextos, com hierarquia de pesos estável;")
add_bullet("O esquema de cores se repete em todos os componentes: botões primários sempre roxos, ações destrutivas sempre vermelhas e feedback de sucesso sempre verde;")
add_bullet("Os espaçamentos obedecem a um grid de 8px (8px, 16px, 24px, 32px), produzindo ritmo visual uniforme;")
add_bullet("Os componentes reutilizáveis (cards, modais, toasts e formulários) possuem aparência e comportamento idênticos em todas as páginas.")
add_par(
    "Essa consistência responde à Heurística 4 de Nielsen (1994) — consistência e padrões "
    "—, que postula que usuários não deveriam precisar se perguntar se palavras, situações "
    "ou ações diferentes significam a mesma coisa.",
    primeira_linha=1.25
)

add_heading("5.3.2 Organização lógica e hierarquia visual", 3)
add_par(
    "A reformulação da home (Seção 4.9) implementa uma hierarquia visual em três níveis:",
    primeira_linha=1.25
)
add_bullet("Nível 1 — destaque máximo: carrossel de destaques no topo, com imagens grandes ocupando 100% da largura;")
add_bullet("Nível 2 — ações principais: cards grandes de Comparar Preços, Minha Lista e Mercados Próximos, posicionados logo abaixo do carrossel;")
add_bullet("Nível 3 — ações secundárias: cards menores de Ranking, Mapa e demais funcionalidades, dispostos na sequência, ainda disponíveis mas sem competir pela atenção inicial.")
add_par(
    "Essa organização dialoga com a Lei de Fitts e com o princípio de proximidade da "
    "Gestalt: elementos relacionados ficam visualmente próximos, e as ações mais relevantes "
    "são maiores e estão mais perto do ponto inicial de atenção do usuário.",
    primeira_linha=1.25
)

add_heading("5.3.3 Padrões de design reconhecíveis", 3)
add_par(
    "As melhorias adotam convenções de UI já consolidadas no repertório dos usuários:",
    primeira_linha=1.25
)
add_bullet("Login em layout split-screen, com formulário à esquerda e imagem à direita, padrão comum em aplicações web modernas;")
add_bullet("Menu hamburger para navegação em telas móveis, convenção universal em interfaces compactas;")
add_bullet("Sistema de estrelas de 1 a 5 para avaliações, padrão estabelecido em plataformas como Google Maps e Amazon;")
add_bullet("Campo de busca com ícone de lupa no canto superior, em posição e iconografia familiares aos usuários da web;")
add_bullet("Botão de ação flutuante (FAB) para ação primária no mobile, convenção consagrada pelo Material Design.")

# 5.4 QUALIDADE DE USO
add_heading("5.4 Qualidade de Uso", 2)
add_heading("5.4.1 Eficiência", 3)
add_par(
    "Eficiência refere-se ao número mínimo de ações necessário para completar uma tarefa. "
    "As melhorias reduzem significativamente o esforço do usuário:",
    primeira_linha=1.25
)
add_bullet("O campo de preço com máscara automática (Seção 4.7) faz com que o usuário digite apenas os números (“1599”), enquanto o sistema formata para “R$ 15,99”. A tarefa, que envolvia mais de seis ações (incluindo digitar “R$”, vírgula e posicionar o cursor), reduz-se a quatro dígitos;")
add_bullet("A geolocalização automática no mapa (Seção 4.4) elimina a etapa manual de procurar a própria localização, economizando três ou quatro interações (abrir busca, digitar cidade, selecionar, confirmar);")
add_bullet("O leitor de código de barras pela câmera (Seção 4.6) permite que apontar para o produto preencha sozinho nome, marca e categoria, suprimindo a digitação de três campos;")
add_bullet("O botão “Ver detalhes” no popup do mapa (Seção 4.4) reduz a chegada à página do mercado a um único clique, contra os três ou mais necessários no fluxo de navegação tradicional.")

add_heading("5.4.2 Eficácia", 3)
add_par(
    "Eficácia diz respeito à capacidade de o usuário concluir tarefas corretamente. As "
    "melhorias retiram ambiguidades e fontes de erro:",
    primeira_linha=1.25
)
add_bullet("A separação de perfis (Seção 4.5) elimina o risco de consumidores cadastrarem preços por engano. A tarefa “informar preço” passa a ser executada apenas por quem tem legitimidade (comerciante verificado);")
add_bullet("A validação por faixa de preço (Seção 4.6) interrompe o fluxo quando um valor incoerente é informado, dando ao usuário a chance de corrigir antes de gravar dados inválidos;")
add_bullet("A máscara monetária, aceitando vírgula e ponto (Seção 4.7), elimina os erros de formato que antes provocavam falhas no salvamento.")

add_heading("5.4.3 Satisfação", 3)
add_par(
    "Satisfação trata da experiência ser fluida, agradável e livre de frustrações:",
    primeira_linha=1.25
)
add_bullet("A tipografia legível (Seção 4.1) resolve a queixa relacionada à leitura, sobretudo entre usuários com baixa visão;")
add_bullet("O contraste adequado (Seção 4.2) elimina o desconforto de “não conseguir enxergar” o texto, tornando a leitura confortável em diferentes condições de iluminação;")
add_bullet("A home redesenhada (Seção 4.9) substitui a sensação de tela vazia por uma proposta visual orientada e hierarquizada;")
add_bullet("Os feedbacks visuais a cada ação (toasts, spinners e animações) reduzem a ansiedade do usuário e reforçam sua confiança na plataforma.")

add_heading("5.4.4 Acessibilidade", 3)
add_par(
    "As ações voltadas à acessibilidade seguem as diretrizes da WCAG 2.1, com foco nos "
    "níveis AA e AAA. Esse foco se justifica pelo perfil heterogêneo do público-alvo, que "
    "inclui usuários idosos e pessoas com limitações visuais:",
    primeira_linha=1.25
)
add_bullet("Contraste de cores (Seção 4.2): razão mínima de 7:1 para texto normal e 4,5:1 para texto grande, atendendo ao nível AAA da WCAG 2.1, critério 1.4.6;")
add_bullet("Responsividade completa (Seção 4.10): layout adaptável de 320px a 1920px, em conformidade com o critério 1.4.10 (Reflow), que assegura acessibilidade do conteúdo sem rolagem horizontal em telas pequenas;")
add_bullet("Navegação por teclado: todos os elementos interativos respondem a Tab e Shift+Tab, com indicadores visuais de foco (outline roxo de 2px), em atendimento ao critério 2.1.1 (Keyboard);")
add_bullet("Textos alternativos em todas as imagens de produtos e em capturas de tela, conforme o critério 1.1.1 (Non-text Content);")
add_bullet("Atributos ARIA nos componentes dinâmicos — modais (role=“dialog”, aria-modal), toasts (role=“alert”) e menus (aria-expanded) —, garantindo compatibilidade com leitores de tela.")

add_pagebreak()

# ==========================================================
# 6 QUADRO RESUMO
# ==========================================================
add_heading("6 QUADRO RESUMO – MELHORIAS E CONCEITOS APLICADOS", 1)
add_par(
    "O quadro a seguir sintetiza a relação entre cada melhoria documentada na Parte 1 e os "
    "conceitos-chave de IHC aplicados na Parte 2. A leitura do quadro confirma que as "
    "quatro dimensões — interação, affordance, interface e qualidade de uso — foram "
    "contempladas ao longo do projeto, evitando o predomínio isolado de um único aspecto.",
    primeira_linha=1.25
)

dados_quadro = [
    ("Melhoria", "Interação", "Affordance", "Interface", "Qualidade de Uso"),
    ("4.1 Tipografia e legibilidade", "", "", "✓", "Satisfação, Acessibilidade"),
    ("4.2 Contraste e cores", "", "✓", "✓", "Satisfação, Acessibilidade"),
    ("4.3 Navegação e botão Voltar", "✓", "✓", "✓", "Eficiência"),
    ("4.4 Mapa interativo", "✓", "✓", "", "Eficiência"),
    ("4.5 Separação de perfis", "✓", "✓", "✓", "Eficácia"),
    ("4.6 Validação anti-fraude", "✓", "", "", "Eficácia"),
    ("4.7 UX do formulário", "✓", "✓", "", "Eficiência, Eficácia"),
    ("4.8 Lista de compras", "✓", "✓", "", "Eficiência"),
    ("4.9 Home redesign", "", "✓", "✓", "Satisfação"),
    ("4.10 Responsividade", "", "", "✓", "Acessibilidade"),
    ("4.11 Avaliações e ranking", "✓", "", "✓", "Eficácia"),
    ("4.12 Correções de bugs", "✓", "", "", "Eficácia"),
]

tabela = doc.add_table(rows=len(dados_quadro), cols=5)
tabela.style = 'Light Grid Accent 1'
tabela.alignment = WD_ALIGN_PARAGRAPH.CENTER

larguras_cm = [5.5, 2.0, 2.0, 2.0, 4.0]
for j, larg in enumerate(larguras_cm):
    for i in range(len(dados_quadro)):
        tabela.cell(i, j).width = Cm(larg)

for i, linha in enumerate(dados_quadro):
    for j, conteudo in enumerate(linha):
        cell = tabela.cell(i, j)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_cell = cell.paragraphs[0]
        p_cell.paragraph_format.space_after = Pt(0)
        p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run_cell = p_cell.add_run(conteudo)
        run_cell.font.name = 'Times New Roman'
        run_cell.font.size = Pt(10)
        if i == 0:
            run_cell.bold = True

add_par(" ", justificar=False, espaco_depois=2)
add_par(
    "Pela leitura do quadro, observa-se que as melhorias de maior abrangência conceitual "
    "(linhas 4.3, 4.5 e 4.7) atendem simultaneamente a três ou quatro conceitos, sinalizando "
    "que são os pontos de maior impacto na experiência do usuário. Por outro lado, ajustes "
    "como tipografia e responsividade, embora concentrados em poucas dimensões, são "
    "decisivos para o atendimento dos critérios de acessibilidade e satisfação.",
    primeira_linha=1.25
)
add_pagebreak()

# ==========================================================
# 7 ANÁLISE CRÍTICA E CONSIDERAÇÕES FINAIS
# ==========================================================
add_heading("7 ANÁLISE CRÍTICA E CONSIDERAÇÕES FINAIS", 1)
add_par(
    "Olhando para o conjunto das melhorias implementadas, fica claro que o ganho mais "
    "expressivo desta etapa do projeto foi a transição de uma interface centrada na entrega "
    "funcional para uma interface centrada no usuário. Na primeira entrega, o foco "
    "predominante era “a aplicação funciona”. A pergunta orientadora desta segunda "
    "fase passou a ser “a aplicação funciona bem para quem vai usá-la, no contexto real "
    "de uso?”.",
    primeira_linha=1.25
)
add_par(
    "Esse deslocamento de perspectiva, embora pareça sutil, teve consequências práticas "
    "significativas. A separação de perfis, por exemplo, deixou de ser apenas uma exigência "
    "técnica do modelo de dados e passou a ser entendida como uma decisão de design que "
    "afeta a confiabilidade da informação apresentada ao consumidor. Da mesma forma, a "
    "tipografia e o contraste deixaram de ser questões de gosto estético para tornarem-se "
    "decisões de inclusão e acessibilidade.",
    primeira_linha=1.25
)
add_par(
    "Entre as escolhas mais relevantes, destacam-se três. A primeira foi a aplicação "
    "rigorosa das diretrizes de contraste da WCAG 2.1, garantindo legibilidade ao "
    "público-alvo da plataforma. A segunda foi a reformulação do fluxo de cadastro de "
    "produtos, com máscara monetária, validação por faixa de preço e separação de perfis, "
    "que em conjunto mitigam o risco de dados fraudulentos sem onerar o usuário legítimo. A "
    "terceira foi a unificação do feedback visual em toda a aplicação, com toasts, spinners "
    "e animações coerentes, transformando a experiência em um diálogo previsível.",
    primeira_linha=1.25
)
add_par(
    "É preciso reconhecer, ao mesmo tempo, as limitações do processo. O feedback coletado "
    "veio principalmente do professor e do próprio grupo, sem a realização de testes "
    "formais com usuários externos. Em uma próxima iteração, seria importante conduzir "
    "testes de usabilidade com representantes do público-alvo — especialmente pessoas mais "
    "velhas, que já foram identificadas como um perfil estratégico para o MercadoCerto. A "
    "ausência desse retorno direto torna algumas decisões mais dependentes de inferência do "
    "que de evidência.",
    primeira_linha=1.25
)
add_par(
    "Outra limitação relevante é a validação automática de preços por OCR, que permanece "
    "como uma melhoria prevista, mas ainda não implementada. A versão atual depende da "
    "obrigatoriedade da foto da etiqueta como prova visual, o que reduz riscos, mas não "
    "elimina por completo a possibilidade de inconsistência. Trata-se de um ponto natural "
    "de evolução para as próximas entregas.",
    primeira_linha=1.25
)
add_par(
    "No conjunto, a experiência de revisitar o projeto sob a ótica explícita dos conceitos "
    "de IHC foi a contribuição formativa mais importante deste trabalho. Mais do que "
    "implementar melhorias, foi possível justificar por que cada uma delas faz sentido — "
    "exercício que aproxima a prática de desenvolvimento de uma postura reflexiva sobre o "
    "papel do projetista de interfaces. Os autores compreendem que a interface ora entregue "
    "não é um produto final, mas uma versão intermediária de um processo contínuo de "
    "avaliação, aprendizado e refinamento.",
    primeira_linha=1.25
)

add_pagebreak()

# ==========================================================
# REFERÊNCIAS
# ==========================================================
add_heading("REFERÊNCIAS", 1)

referencias = [
    "COOPER, A.; REIMANN, R.; CRONIN, D.; NOESSEL, C. About Face: The Essentials of Interaction Design. 4. ed. Indianapolis: Wiley, 2014.",
    "GIBSON, J. J. The Ecological Approach to Visual Perception. Boston: Houghton Mifflin, 1979.",
    "INTERNATIONAL ORGANIZATION FOR STANDARDIZATION. ISO 9241-11: Ergonomics of human-system interaction — Part 11: Usability: Definitions and concepts. Geneva: ISO, 2018.",
    "KRUG, S. Don't Make Me Think, Revisited: A Common Sense Approach to Web Usability. 3. ed. Berkeley: New Riders, 2014.",
    "MAYHEW, D. J. Principles and Guidelines in Software User Interface Design. Englewood Cliffs: Prentice Hall, 1992.",
    "NIELSEN, J. Usability Engineering. San Diego: Academic Press, 1993.",
    "NIELSEN, J. 10 Usability Heuristics for User Interface Design. Nielsen Norman Group, 1994. Disponível em: https://www.nngroup.com/articles/ten-usability-heuristics/. Acesso em: 15 mai. 2026.",
    "NORMAN, D. A. The Design of Everyday Things. Revised and Expanded Edition. New York: Basic Books, 2013.",
    "PEDROSO JUNIOR, C. A. UX – Arquitetura e Usabilidade: Aula VIII (08/05/2025). Centro Universitário Senac, 2025. Material didático.",
    "PEDROSO JUNIOR, C. A. UX – Arquitetura e Usabilidade: Aula IX (15/05/2026). Centro Universitário Senac, 2026. Material didático.",
    "PREECE, J.; ROGERS, Y.; SHARP, H. Design de Interação: além da interação humano-computador. 3. ed. Porto Alegre: Bookman, 2013.",
    "WINOGRAD, T. From computing machinery to interaction design. In: DENNING, P.; METCALFE, R. (eds.). Beyond Calculation: The Next Fifty Years of Computing. New York: Springer-Verlag, 2003.",
    "W3C. Web Content Accessibility Guidelines (WCAG) 2.1. World Wide Web Consortium, 2018. Disponível em: https://www.w3.org/TR/WCAG21/. Acesso em: 15 mai. 2026.",
]

for ref in referencias:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ==========================================================
# Salvar
# ==========================================================
doc.save(PATH_OUT)
print("Documento gerado com sucesso em:", PATH_OUT)

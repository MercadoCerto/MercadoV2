from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import re

doc = Document()

# ── Estilos ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x4b, 0x00, 0x82)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(14)
        hs.paragraph_format.space_after = Pt(6)

# ── Helpers ──────────────────────────────────────────────
def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

# ═══════════════════════════════════════════════════════════
#  CAPA
# ═══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MercadoCerto')
r.font.size = Pt(42)
r.font.color.rgb = RGBColor(0x4b, 0x00, 0x82)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Documentacao Completa do Sistema')
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Plataforma de comparacao de precos e avaliacao de mercados')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Versao 1.0 — Marco 2026')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  SUMARIO
# ═══════════════════════════════════════════════════════════
doc.add_heading('Sumario', level=1)
sumario = [
    '1. Visao Geral',
    '2. Tecnologias',
    '3. Arquitetura',
    '4. Configuracao e Instalacao',
    '5. Banco de Dados',
    '6. API REST — Endpoints',
    '7. Models (Entidades JPA)',
    '8. Services (Logica de Negocio)',
    '9. DTOs (Data Transfer Objects)',
    '10. Frontend — Paginas',
    '11. Autenticacao e Seguranca',
    '12. Funcionalidades Principais',
    '13. Estrutura de Arquivos',
]
for item in sumario:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  1. VISAO GERAL
# ═══════════════════════════════════════════════════════════
doc.add_heading('1. Visao Geral', level=1)
doc.add_paragraph(
    'MercadoCerto e uma aplicacao web fullstack para comparacao de precos entre mercados e supermercados. '
    'O sistema conecta consumidores e comerciantes em uma plataforma unica onde e possivel comparar precos, '
    'avaliar estabelecimentos, criar listas de compras inteligentes e localizar mercados no mapa.'
)

doc.add_heading('Para Consumidores', level=3)
add_bullet('Comparar precos de produtos entre diferentes mercados')
add_bullet('Avaliar e dar notas a estabelecimentos')
add_bullet('Criar listas de compras com otimizacao automatica de precos')
add_bullet('Localizar mercados proximos no mapa interativo')
add_bullet('Acompanhar historico de precos ao longo do tempo')

doc.add_heading('Para Comerciantes', level=3)
add_bullet('Cadastrar mercados e produtos com imagens')
add_bullet('Publicar e atualizar precos em tempo real')
add_bullet('Acompanhar avaliacoes e metricas de reputacao')
add_bullet('Visualizar dashboard com evolucao mensal de notas')

doc.add_paragraph(
    'O backend e uma API REST construida com Spring Boot 3.3.5 (Java 17) e o frontend e composto por '
    'paginas HTML estaticas com JavaScript puro (sem frameworks), servidas pelo proprio Spring Boot.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  2. TECNOLOGIAS
# ═══════════════════════════════════════════════════════════
doc.add_heading('2. Tecnologias', level=1)

doc.add_heading('Backend', level=2)
add_table(
    ['Tecnologia', 'Versao', 'Funcao'],
    [
        ['Java', '17', 'Linguagem principal'],
        ['Spring Boot', '3.3.5', 'Framework web'],
        ['Spring Data JPA', '3.3.x', 'ORM / Acesso a dados'],
        ['Spring Security', '6.x', 'Autenticacao (BCrypt)'],
        ['MariaDB', '10.x+', 'Banco de dados relacional'],
        ['Lombok', '1.18.x', 'Reducao de boilerplate'],
        ['Maven', '3.9.x', 'Gerenciamento de dependencias'],
    ]
)

doc.add_heading('Frontend', level=2)
add_table(
    ['Tecnologia', 'Funcao'],
    [
        ['HTML5 / CSS3 / JavaScript (ES6+)', 'Interface do usuario'],
        ['Chart.js 4.4.1', 'Graficos (historico de precos, reputacao)'],
        ['Leaflet.js 1.9.4', 'Mapa interativo (OpenStreetMap)'],
        ['Font Awesome 6.5.0', 'Icones (pagina inicial)'],
        ['Google Fonts (Syne + DM Sans)', 'Tipografia'],
    ]
)

doc.add_heading('Design', level=2)
add_bullet('Roxo (#4b0082) como cor primaria', 'Tema: ')
add_bullet('Syne (titulos) e DM Sans (corpo)', 'Fontes: ')
add_bullet('Responsivo com CSS Grid e Flexbox', 'Layout: ')
add_bullet('Cards, modais, toasts, spinners, tabs', 'Componentes: ')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  3. ARQUITETURA
# ═══════════════════════════════════════════════════════════
doc.add_heading('3. Arquitetura', level=1)
doc.add_paragraph('O projeto segue a arquitetura em camadas padrao do Spring Boot:')

add_code(
    '[Frontend HTML/JS]\n'
    '        |\n'
    '        | HTTP (fetch API)\n'
    '        v\n'
    '[Controllers]  -->  Recebem requisicoes HTTP, retornam JSON\n'
    '        |\n'
    '        v\n'
    '[Services]  ----->  Logica de negocio, validacoes, calculos\n'
    '        |\n'
    '        v\n'
    '[Repositories]  ->  Acesso ao banco via Spring Data JPA\n'
    '        |\n'
    '        v\n'
    '[MariaDB]  ------>  Persistencia de dados'
)

doc.add_heading('Fluxo de Autenticacao', level=2)
doc.add_paragraph(
    'A autenticacao e feita no frontend via localStorage. O backend valida credenciais no login e retorna '
    'os dados do usuario. Nao ha JWT nem sessoes server-side — o estado de login e mantido exclusivamente no navegador.'
)
add_code(
    '1. POST /api/usuarios/login\n'
    '2. Backend valida senha (BCrypt)\n'
    '3. Retorna LoginResponseDTO\n'
    '4. Frontend salva em localStorage("mc_usuario")\n'
    '5. Paginas protegidas verificam localStorage no carregamento'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  4. CONFIGURACAO E INSTALACAO
# ═══════════════════════════════════════════════════════════
doc.add_heading('4. Configuracao e Instalacao', level=1)

doc.add_heading('Pre-requisitos', level=2)
add_bullet('Java 17+')
add_bullet('Maven 3.9+')
add_bullet('MariaDB 10.x+ (rodando na porta 3306)')

doc.add_heading('Passo a passo', level=2)

p = doc.add_paragraph()
r = p.add_run('1. Criar o banco de dados:')
r.bold = True
add_code('CREATE DATABASE MercadoCerto;')

p = doc.add_paragraph()
r = p.add_run('2. Executar o script de criacao (opcional):')
r.bold = True
add_code('mysql -u root -p MercadoCerto < mercdadocertocreate.sql')

p = doc.add_paragraph()
r = p.add_run('3. Configurar credenciais em application.properties:')
r.bold = True
add_code(
    'spring.datasource.url=jdbc:mariadb://localhost:3306/MercadoCerto\n'
    'spring.datasource.username=root\n'
    'spring.datasource.password=root'
)

p = doc.add_paragraph()
r = p.add_run('4. Compilar e executar:')
r.bold = True
add_code('mvn clean compile\nmvn spring-boot:run')

p = doc.add_paragraph()
r = p.add_run('5. Acessar: ')
r.bold = True
p.add_run('http://localhost:8080')

doc.add_heading('Propriedades de Configuracao', level=2)
add_table(
    ['Propriedade', 'Valor', 'Descricao'],
    [
        ['server.port', '8080', 'Porta do servidor'],
        ['spring.jpa.hibernate.ddl-auto', 'update', 'Auto-criacao de tabelas'],
        ['spring.jpa.show-sql', 'true', 'Exibir queries no console'],
        ['spring.servlet.multipart.max-file-size', '10MB', 'Tamanho maximo de upload'],
        ['uploads.path', 'uploads', 'Diretorio para imagens'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  5. BANCO DE DADOS
# ═══════════════════════════════════════════════════════════
doc.add_heading('5. Banco de Dados', level=1)

doc.add_heading('Diagrama ER (simplificado)', level=2)
add_code(
    'usuario (1) ---- (N) lista_compras (1) ---- (N) item_lista\n'
    '    |                                              |\n'
    '    | (0..N)                                  id_produto\n'
    '    v                                              v\n'
    'avaliacao (N) ----------- (1) mercado         produto\n'
    '                               |                |\n'
    '                               | (N)            | (N)\n'
    '                               v                v\n'
    '                             preco <---------- preco\n'
    '                    (id_mercado + id_produto + valor + data_hora)'
)

doc.add_heading('Tabela: usuario', level=2)
add_table(
    ['Coluna', 'Tipo', 'Restricoes', 'Descricao'],
    [
        ['id_usuario', 'INT', 'PK, AUTO_INCREMENT', 'Identificador'],
        ['nome_usuario', 'VARCHAR(100)', 'NOT NULL', 'Nome completo'],
        ['email', 'VARCHAR(100)', 'UNIQUE, NOT NULL', 'E-mail'],
        ['login', 'VARCHAR(50)', 'UNIQUE, NOT NULL', 'Nome de usuario'],
        ['senha', 'VARCHAR(255)', 'NOT NULL', 'Hash BCrypt'],
        ['tipo_conta', 'ENUM', 'NOT NULL', 'USUARIO ou COMERCIO'],
        ['cnpj', 'VARCHAR(14)', 'UNIQUE, NULL', 'CNPJ (so comercio)'],
        ['cpf', 'VARCHAR(11)', 'UNIQUE, NULL', 'CPF'],
        ['data_nascimento', 'DATE', 'NULL', 'Data de nascimento'],
        ['telefone', 'VARCHAR(20)', 'NULL', 'Telefone'],
        ['id_endereco', 'INT', 'FK -> endereco', 'Endereco'],
    ]
)

doc.add_heading('Tabela: mercado', level=2)
add_table(
    ['Coluna', 'Tipo', 'Restricoes', 'Descricao'],
    [
        ['id_mercado', 'INT', 'PK, AUTO_INCREMENT', 'Identificador'],
        ['nome_mercado', 'VARCHAR(100)', 'NOT NULL', 'Nome do mercado'],
        ['latitude', 'DOUBLE', 'NULL', 'Coordenada GPS'],
        ['longitude', 'DOUBLE', 'NULL', 'Coordenada GPS'],
    ]
)

doc.add_heading('Tabela: produto', level=2)
add_table(
    ['Coluna', 'Tipo', 'Restricoes', 'Descricao'],
    [
        ['id_produto', 'INT', 'PK, AUTO_INCREMENT', 'Identificador'],
        ['nome_produto', 'VARCHAR(100)', 'NOT NULL', 'Nome do produto'],
        ['marca', 'VARCHAR(100)', 'NULL', 'Marca'],
        ['categoria', 'VARCHAR(50)', 'NULL', 'Categoria'],
        ['codigo_barras', 'VARCHAR(50)', 'NULL', 'Codigo de barras'],
        ['imagem', 'VARCHAR(255)', 'NULL', 'Nome do arquivo de imagem'],
        ['preco', 'DOUBLE', 'NULL', 'Preco base de cadastro'],
        ['validade', 'VARCHAR(20)', 'NULL', 'Data de validade'],
        ['latitude', 'DOUBLE', 'NULL', 'Local de cadastro'],
        ['longitude', 'DOUBLE', 'NULL', 'Local de cadastro'],
    ]
)

doc.add_heading('Tabela: preco', level=2)
add_table(
    ['Coluna', 'Tipo', 'Restricoes', 'Descricao'],
    [
        ['id_preco', 'INT', 'PK, AUTO_INCREMENT', 'Identificador'],
        ['id_produto', 'INT', 'FK -> produto', 'Produto referenciado'],
        ['id_mercado', 'INT', 'FK -> mercado', 'Mercado que pratica o preco'],
        ['valor', 'DECIMAL(10,2)', 'NOT NULL', 'Valor em reais'],
        ['data_hora', 'DATETIME', 'DEFAULT NOW()', 'Data/hora da publicacao'],
    ]
)

doc.add_heading('Tabela: avaliacao', level=2)
add_table(
    ['Coluna', 'Tipo', 'Restricoes', 'Descricao'],
    [
        ['id_avaliacao', 'INT', 'PK, AUTO_INCREMENT', 'Identificador'],
        ['id_mercado', 'INT', 'FK -> mercado', 'Mercado avaliado'],
        ['id_usuario', 'INT', 'FK, NULL', 'Usuario (NULL = anonimo)'],
        ['nota', 'INT', 'CHECK(1..5)', 'Nota de 1 a 5'],
        ['comentario', 'TEXT', 'NULL', 'Comentario opcional'],
        ['data_avaliacao', 'DATETIME', 'DEFAULT NOW()', 'Data/hora'],
    ]
)

doc.add_heading('Tabela: lista_compras', level=2)
add_table(
    ['Coluna', 'Tipo', 'Restricoes', 'Descricao'],
    [
        ['id_lista', 'INT', 'PK, AUTO_INCREMENT', 'Identificador'],
        ['id_usuario', 'INT', 'FK -> usuario', 'Dono da lista'],
        ['nome_lista', 'VARCHAR(100)', 'NOT NULL', 'Nome da lista'],
        ['data_criacao', 'DATETIME', 'DEFAULT NOW()', 'Data de criacao'],
    ]
)

doc.add_heading('Tabela: item_lista', level=2)
add_table(
    ['Coluna', 'Tipo', 'Restricoes', 'Descricao'],
    [
        ['id_item', 'INT', 'PK, AUTO_INCREMENT', 'Identificador'],
        ['id_lista', 'INT', 'FK -> lista_compras (CASCADE)', 'Lista pai'],
        ['id_produto', 'INT', 'NOT NULL', 'Produto referenciado'],
        ['quantidade', 'INT', 'DEFAULT 1', 'Quantidade desejada'],
    ]
)

doc.add_paragraph('Tabelas auxiliares: endereco, cidade e pais formam a hierarquia de localizacao (pais -> cidade -> endereco).')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  6. API REST
# ═══════════════════════════════════════════════════════════
doc.add_heading('6. API REST — Endpoints', level=1)
doc.add_paragraph('Base URL: http://localhost:8080')

# Usuarios
doc.add_heading('Usuarios (/api/usuarios)', level=2)
add_table(
    ['Metodo', 'Endpoint', 'Descricao'],
    [
        ['POST', '/api/usuarios/register', 'Cadastrar usuario'],
        ['POST', '/api/usuarios/login', 'Fazer login'],
        ['GET', '/api/usuarios', 'Listar todos'],
        ['GET', '/api/usuarios/{id}', 'Buscar por ID'],
        ['GET', '/api/usuarios/login/{login}', 'Buscar por login'],
        ['GET', '/api/usuarios/email/{email}', 'Buscar por email'],
    ]
)

p = doc.add_paragraph()
r = p.add_run('Body de registro: ')
r.bold = True
add_code('{ "nomeUsuario", "email", "login", "senha", "tipoConta", "cnpj?" }')

p = doc.add_paragraph()
r = p.add_run('Body de login: ')
r.bold = True
add_code('{ "login", "senha" }')

p = doc.add_paragraph()
r = p.add_run('Resposta de login (LoginResponseDTO):')
r.bold = True
add_code(
    '{\n'
    '  "idUsuario": 1,\n'
    '  "nomeUsuario": "Joao Silva",\n'
    '  "email": "joao@email.com",\n'
    '  "login": "joao",\n'
    '  "tipoConta": "USUARIO"\n'
    '}'
)

# Produtos
doc.add_heading('Produtos (/api/produtos)', level=2)
add_table(
    ['Metodo', 'Endpoint', 'Descricao'],
    [
        ['GET', '/api/produtos', 'Listar todos os produtos'],
        ['POST', '/api/produtos/cadastrar', 'Cadastrar produto (multipart/form-data)'],
    ]
)

p = doc.add_paragraph()
r = p.add_run('Cadastro (multipart): ')
r.bold = True
p.add_run('campo "produto" (JSON), campo "idMercado" (number), campo "imagem" (file, opcional)')

# Mercados
doc.add_heading('Mercados (/api/mercados)', level=2)
add_table(
    ['Metodo', 'Endpoint', 'Descricao'],
    [
        ['GET', '/api/mercados', 'Listar todos'],
        ['GET', '/api/mercados/{id}', 'Buscar por ID'],
        ['POST', '/api/mercados', 'Cadastrar mercado'],
        ['PUT', '/api/mercados/{id}', 'Atualizar mercado'],
        ['DELETE', '/api/mercados/{id}', 'Remover mercado'],
        ['GET', '/api/mercados/proximos', 'Buscar proximos (?latitude, longitude, raioKm)'],
    ]
)

p = doc.add_paragraph()
r = p.add_run('Body: ')
r.bold = True
add_code('{ "nomeMercado", "latitude", "longitude" }')

# Precos
doc.add_heading('Precos (/api/precos)', level=2)
add_table(
    ['Metodo', 'Endpoint', 'Descricao'],
    [
        ['POST', '/api/precos', 'Publicar preco'],
        ['GET', '/api/precos/comparar?produto={id}', 'Comparar precos de produto em mercados'],
        ['GET', '/api/precos/historico?produto={id}&mercado={id}', 'Historico de precos'],
        ['GET', '/api/precos/mercado/{idMercado}', 'Catalogo do mercado'],
    ]
)

p = doc.add_paragraph()
r = p.add_run('Resposta de comparacao (ComparadorDTO):')
r.bold = True
add_code(
    '{\n'
    '  "idProduto": 9,\n'
    '  "nomeProduto": "Arroz Integral",\n'
    '  "precos": [\n'
    '    { "idMercado": 1, "nomeMercado": "Mercado X", "valor": 5.99, "dataHora": "2026-03-27T20:40:36" },\n'
    '    { "idMercado": 3, "nomeMercado": "Condor", "valor": 7.50, "dataHora": "2026-03-28T10:15:00" }\n'
    '  ]\n'
    '}'
)

# Avaliacoes
doc.add_heading('Avaliacoes (/api/avaliacoes)', level=2)
add_table(
    ['Metodo', 'Endpoint', 'Descricao'],
    [
        ['POST', '/api/avaliacoes', 'Enviar avaliacao'],
        ['GET', '/api/avaliacoes/mercado/{idMercado}', 'Avaliacoes de um mercado'],
        ['GET', '/api/avaliacoes/media/{idMercado}', 'Media de notas'],
        ['GET', '/api/avaliacoes/total/{idMercado}', 'Total de avaliacoes'],
        ['GET', '/api/avaliacoes/ranking', 'Ranking de mercados'],
    ]
)

p = doc.add_paragraph()
r = p.add_run('Body de avaliacao: ')
r.bold = True
add_code('{ "idMercado", "nota" (1-5), "comentario"?, "dataAvaliacao" }')

p = doc.add_paragraph()
r = p.add_run('Resposta de ranking:')
r.bold = True
add_code(
    '[\n'
    '  { "id_mercado": 1, "nome_mercado": "Mercado X", "media": 4.5, "votos": 12 },\n'
    '  { "id_mercado": 3, "nome_mercado": "Condor", "media": 4.2, "votos": 8 }\n'
    ']'
)

# Lista de Compras
doc.add_heading('Lista de Compras (/api/lista)', level=2)
add_table(
    ['Metodo', 'Endpoint', 'Descricao'],
    [
        ['GET', '/api/lista/usuario/{idUsuario}', 'Listas do usuario'],
        ['GET', '/api/lista/{idLista}', 'Detalhes da lista'],
        ['POST', '/api/lista?idUsuario={id}&nomeLista={nome}', 'Criar lista'],
        ['POST', '/api/lista/{idLista}/item', 'Adicionar item'],
        ['DELETE', '/api/lista/{idLista}/item/{idProduto}', 'Remover item'],
        ['DELETE', '/api/lista/{idLista}', 'Excluir lista'],
        ['GET', '/api/lista/{idLista}/otimizar', 'Otimizar lista'],
    ]
)

p = doc.add_paragraph()
r = p.add_run('Resposta de otimizacao (ResultadoOtimizadorDTO):')
r.bold = True
add_code(
    '{\n'
    '  "idLista": 1, "nomeLista": "Compras do mes",\n'
    '  "planos": [\n'
    '    {\n'
    '      "idMercado": 1, "nomeMercado": "Mercado X",\n'
    '      "itens": [{ "nomeProduto": "Arroz", "valor": 5.99, "quantidade": 2 }],\n'
    '      "subtotal": 11.98\n'
    '    }\n'
    '  ],\n'
    '  "totalOtimizado": 45.90,\n'
    '  "totalMercadoMaisBarato": 52.30,\n'
    '  "economia": 6.40\n'
    '}'
)

# Painel
doc.add_heading('Painel do Comercio (/api/painel)', level=2)
add_table(
    ['Metodo', 'Endpoint', 'Descricao'],
    [
        ['GET', '/api/painel/dashboard/{idMercado}', 'Dashboard completo do mercado'],
    ]
)

p = doc.add_paragraph()
r = p.add_run('Resposta (DashboardComercioDTO):')
r.bold = True
add_code(
    '{\n'
    '  "mediaGeral": 4.3,\n'
    '  "totalAvaliacoes": 15,\n'
    '  "totalProdutos": 42,\n'
    '  "evolucaoMensal": [\n'
    '    { "mes": "2026-01", "mediaMes": 4.0, "totalAvaliacoes": 5 }\n'
    '  ],\n'
    '  "avaliacoesRecentes": [...]\n'
    '}'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  7. MODELS
# ═══════════════════════════════════════════════════════════
doc.add_heading('7. Models (Entidades JPA)', level=1)

models = [
    ('Usuario', 'usuario', 'idUsuario, nomeUsuario, email, login, senha (BCrypt, write-only), tipoConta (USUARIO/COMERCIO), cnpj, cpf, dataNascimento, telefone, idEndereco'),
    ('Mercado', 'mercado', 'idMercado, nomeMercado, latitude, longitude'),
    ('Produto', 'produto', 'idProduto, nomeProduto, marca, categoria, codigoBarras, imagem, preco, validade, latitude, longitude. Campo idMercado e @Transient (usado apenas no cadastro)'),
    ('Preco', 'preco', 'idPreco, idProduto, idMercado, valor (BigDecimal), dataHora (@PrePersist auto-timestamp)'),
    ('Avaliacao', 'avaliacao', 'idAvaliacao, idMercado, idUsuario (nullable, anonimo), nota (1-5), comentario, dataAvaliacao'),
    ('ListaCompras', 'lista_compras', 'idLista, idUsuario, nomeLista, dataCriacao. @OneToMany com ItemLista (cascade ALL, orphanRemoval)'),
    ('ItemLista', 'item_lista', 'idItem, lista (@ManyToOne), idProduto, quantidade'),
]

for name, table, fields in models:
    doc.add_heading(name, level=2)
    p = doc.add_paragraph()
    r = p.add_run('Tabela: ')
    r.bold = True
    p.add_run(table)
    p = doc.add_paragraph()
    r = p.add_run('Campos: ')
    r.bold = True
    p.add_run(fields)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  8. SERVICES
# ═══════════════════════════════════════════════════════════
doc.add_heading('8. Services (Logica de Negocio)', level=1)

doc.add_heading('UsuarioService', level=2)
add_bullet('Valida email e login unicos')
add_bullet('Valida CNPJ com algoritmo de digitos verificadores (para contas COMERCIO)')
add_bullet('Criptografa senha com BCrypt (strength 10)')
add_bullet('Login aceita tanto login quanto email como identificador')

doc.add_heading('ProdutoService', level=2)
add_bullet('Salva imagens no disco (uploads/) com nome: timestamp_nomeOriginal')
add_bullet('Ao cadastrar produto com idMercado, cria automaticamente registro na tabela preco')

doc.add_heading('PrecoService', level=2)
add_bullet('Comparacao: busca preco atual de um produto em todos os mercados, retorna ordenado por valor')
add_bullet('Historico: retorna evolucao de precos ao longo do tempo')
add_bullet('Catalogo: lista todos os precos atuais de um mercado')

doc.add_heading('AvaliacaoService', level=2)
add_bullet('Aceita avaliacoes anonimas (idUsuario nullable)')
add_bullet('Auto-timestamp na data da avaliacao')
add_bullet('Ranking via query nativa: agrupa por mercado, calcula media e votos, ordena por media')

doc.add_heading('ListaComprasService', level=2)
add_bullet('CRUD completo de listas e itens (com verificacao de duplicatas)')
add_bullet('Otimizador: busca melhor preco de cada item, distribui entre mercados, calcula economia vs mercado unico mais barato', 'Algoritmo: ')

doc.add_heading('MercadoService', level=2)
add_bullet('CRUD completo de mercados')
add_bullet('Busca por proximidade usando formula de Haversine (raio em km)')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  9. DTOs
# ═══════════════════════════════════════════════════════════
doc.add_heading('9. DTOs (Data Transfer Objects)', level=1)
doc.add_paragraph('Todos os DTOs sao Java Records (imutaveis), definidos em DTOs.java:')

add_table(
    ['DTO', 'Usado em', 'Campos'],
    [
        ['LoginResponseDTO', 'Login', 'idUsuario, nomeUsuario, email, login, tipoConta'],
        ['ComparadorDTO', 'Comparar precos', 'idProduto, nomeProduto, List<PrecoMercadoDTO>'],
        ['PrecoMercadoDTO', 'Comparar precos', 'idMercado, nomeMercado, valor, dataHora'],
        ['ResultadoOtimizadorDTO', 'Otimizar lista', 'idLista, nomeLista, planos, totalOtimizado, totalMercadoMaisBarato, economia'],
        ['PlanoMercadoDTO', 'Otimizar lista', 'idMercado, nomeMercado, lat, lng, itens, subtotal'],
        ['ItemOtimizadoDTO', 'Otimizar lista', 'idProduto, nomeProduto, idMercado, nomeMercado, valor, quantidade'],
        ['DashboardComercioDTO', 'Painel comercio', 'mediaGeral, totalAvaliacoes, totalProdutos, evolucaoMensal, avaliacoesRecentes'],
        ['ReputacaoMensalDTO', 'Painel comercio', 'mes, mediaMes, totalAvaliacoes'],
        ['ItemListaRequestDTO', 'Adicionar item', 'idProduto, quantidade'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  10. FRONTEND
# ═══════════════════════════════════════════════════════════
doc.add_heading('10. Frontend — Paginas', level=1)

doc.add_heading('Paginas Publicas', level=2)
add_table(
    ['Pagina', 'Arquivo', 'Descricao'],
    [
        ['Home', 'index.html', 'Landing page com grid de acoes rapidas'],
        ['Login', 'login.html', 'Formulario split-screen, redireciona logados'],
        ['Cadastro', 'cadastro.html', 'Toggle Consumidor/Supermercado, validacao CNPJ'],
        ['Produtos', 'lista-produtos.html', 'Grid filtros/ordenacao, botao comparar em cada card'],
        ['Mercados', 'mercados.html', 'Lista com avaliacoes, modal de avaliacao com estrelas'],
        ['Ranking', 'ranking.html', 'Ranking por nota, top 3 com medalhas'],
        ['Mapa', 'mapa.html', 'Leaflet/OpenStreetMap, busca produtos no mapa'],
        ['Comparar', 'comparar.html', 'Tabela de precos + grafico historico (Chart.js)'],
        ['Lista de Compras', 'lista-compra.html', 'Listas via localStorage, checklist interativo'],
    ]
)

doc.add_heading('Paginas Protegidas (requer login)', level=2)
add_table(
    ['Pagina', 'Arquivo', 'Requer', 'Descricao'],
    [
        ['Dashboard', 'dashboard.html', 'USUARIO', 'Stats, listas, top mercados, produtos baratos'],
        ['Painel Comercio', 'painel-comercio.html', 'COMERCIO', '4 abas: reputacao, catalogo, publicar, avaliacoes'],
        ['Cadastrar Produto', 'cadastro-produto.html', 'COMERCIO', 'Upload imagem, selecao mercado'],
        ['Cadastrar Mercado', 'cadastro-mercado.html', 'COMERCIO', 'Geolocalizacao automatica'],
    ]
)

doc.add_heading('Arquivos JavaScript', level=2)
add_table(
    ['Arquivo', 'Descricao'],
    [
        ['js/main.js', 'Constante API, helpers auth, toast, validacao (email, CNPJ), spinner, estrelas, nav'],
        ['js/header.js', 'Injecao dinamica do header baseado no estado de login e tipo de conta'],
    ]
)

doc.add_heading('CSS', level=2)
doc.add_paragraph('Arquivo unico css/style.css com design system completo: tokens de cor, tipografia, grid responsivo, cards, formularios, botoes, badges, toasts, modais, estrelas, rankings, paginas de autenticacao.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  11. AUTENTICACAO
# ═══════════════════════════════════════════════════════════
doc.add_heading('11. Autenticacao e Seguranca', level=1)

doc.add_heading('Backend', level=2)
add_bullet('BCrypt (strength 10) para hash de senhas')
add_bullet('CSRF desabilitado (API REST stateless)')
add_bullet('CORS configurado para localhost:5500 e 127.0.0.1:5500')
add_bullet('Todos os endpoints sao publicos (sem autorizacao server-side)')
add_bullet('Upload seguro: nomes de arquivo sanitizados com timestamp')
add_bullet('Tratamento global de excecoes (GlobalExceptionHandler): 400, 413, 500')

doc.add_heading('Frontend', level=2)
add_bullet('localStorage("mc_usuario") armazena dados do usuario logado')
add_bullet('requireLogin() e requireComercio() verificam e redirecionam imediatamente')
add_bullet('Senhas nunca retornadas pela API (@JsonProperty WRITE_ONLY)')
add_bullet('Validacao completa de CNPJ com digitos verificadores')
add_bullet('Validacao de upload: tipo (image/*) e tamanho (max 5MB) no frontend')

doc.add_heading('Fluxo Completo', level=2)
add_code(
    '1. Usuario acessa /login.html\n'
    '2. Preenche login (ou email) + senha\n'
    '3. POST /api/usuarios/login\n'
    '4. Backend busca por login OU email\n'
    '5. Verifica senha com BCrypt\n'
    '6. Retorna LoginResponseDTO (sem senha)\n'
    '7. Frontend salva em localStorage\n'
    '8. Header se adapta (nome, avatar, dropdown)\n'
    '9. Paginas protegidas verificam localStorage'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  12. FUNCIONALIDADES
# ═══════════════════════════════════════════════════════════
doc.add_heading('12. Funcionalidades Principais', level=1)

doc.add_heading('12.1 Comparacao de Precos', level=2)
doc.add_paragraph(
    'O usuario seleciona um produto e visualiza uma tabela com precos em todos os mercados, '
    'ordenados do mais barato ao mais caro. Inclui badge "Mais barato", diferenca percentual, '
    'e grafico de historico de precos por mercado com filtro de periodo (7/30/90 dias). '
    'Link direto de lista-produtos.html via parametro ?produto=ID na URL.'
)

doc.add_heading('12.2 Avaliacao de Mercados', level=2)
doc.add_paragraph(
    'Qualquer usuario pode avaliar mercados com nota de 1 a 5 estrelas e comentario opcional. '
    'Aceita avaliacoes anonimas (sem login). O sistema calcula media, total de votos e gera '
    'ranking geral ordenado por media com medalhas para os top 3.'
)

doc.add_heading('12.3 Otimizador de Lista de Compras', level=2)
doc.add_paragraph(
    'Dado uma lista de compras, o algoritmo: (1) busca o melhor preco de cada item em todos os mercados, '
    '(2) distribui os itens entre mercados para minimizar custo total, (3) gera planos de compra por mercado '
    'com subtotais, (4) calcula economia comparando com comprar tudo em um unico mercado.'
)

doc.add_heading('12.4 Mapa Interativo', level=2)
doc.add_paragraph(
    'Mapa com Leaflet.js e OpenStreetMap. Busca produtos por nome e exibe marcadores nos mercados '
    'que vendem. Popup com nome, preco, marca e imagem. Botao "Ver Mercados" mostra todos no mapa. '
    'Auto-zoom para enquadrar marcadores.'
)

doc.add_heading('12.5 Painel do Comercio', level=2)
doc.add_paragraph(
    'Dashboard exclusivo para contas COMERCIO com 4 abas: Reputacao (grafico de barras mensal), '
    'Catalogo (precos atuais), Publicar Preco (formulario), Avaliacoes (lista de avaliacoes recebidas).'
)

doc.add_heading('12.6 Dashboard do Consumidor', level=2)
doc.add_paragraph(
    'Painel para contas USUARIO com cards de stats (produtos, mercados, listas, melhor mercado), '
    'listas de compras recentes, top mercados por avaliacao e produtos com menores precos.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
#  13. ESTRUTURA DE ARQUIVOS
# ═══════════════════════════════════════════════════════════
doc.add_heading('13. Estrutura de Arquivos', level=1)

add_code(
    'startup/\n'
    '|-- pom.xml\n'
    '|-- mercdadocertocreate.sql\n'
    '|-- migrar_banco.sql\n'
    '|-- uploads/\n'
    '|\n'
    '|-- src/main/\n'
    '    |-- java/com/mercadocerto/\n'
    '    |   |-- MercadocertoApplication.java\n'
    '    |   |-- config/\n'
    '    |   |   |-- GlobalExceptionHandler.java\n'
    '    |   |   |-- SecurityConfig.java\n'
    '    |   |   |-- WebConfig.java\n'
    '    |   |-- controller/\n'
    '    |   |   |-- AvaliacaoController.java\n'
    '    |   |   |-- DashboardController.java\n'
    '    |   |   |-- ListaComprasController.java\n'
    '    |   |   |-- MercadoController.java\n'
    '    |   |   |-- PrecoController.java\n'
    '    |   |   |-- ProdutoController.java\n'
    '    |   |   |-- UsuarioController.java\n'
    '    |   |-- dto/\n'
    '    |   |   |-- DTOs.java\n'
    '    |   |-- model/\n'
    '    |   |   |-- Avaliacao.java\n'
    '    |   |   |-- ItemLista.java\n'
    '    |   |   |-- ListaCompras.java\n'
    '    |   |   |-- Mercado.java\n'
    '    |   |   |-- Preco.java\n'
    '    |   |   |-- Produto.java\n'
    '    |   |   |-- Usuario.java\n'
    '    |   |-- repository/\n'
    '    |   |   |-- AvaliacaoRepository.java\n'
    '    |   |   |-- ListaComprasRepository.java\n'
    '    |   |   |-- MercadoRepository.java\n'
    '    |   |   |-- PrecoRepository.java\n'
    '    |   |   |-- ProdutoRepository.java\n'
    '    |   |   |-- UsuarioRepository.java\n'
    '    |   |-- service/\n'
    '    |       |-- AvaliacaoService.java\n'
    '    |       |-- ListaComprasService.java\n'
    '    |       |-- MercadoService.java\n'
    '    |       |-- PrecoService.java\n'
    '    |       |-- ProdutoService.java\n'
    '    |       |-- UsuarioService.java\n'
    '    |\n'
    '    |-- resources/\n'
    '        |-- application.properties\n'
    '        |-- static/\n'
    '            |-- css/style.css\n'
    '            |-- js/main.js\n'
    '            |-- js/header.js\n'
    '            |-- index.html\n'
    '            |-- login.html\n'
    '            |-- cadastro.html\n'
    '            |-- dashboard.html\n'
    '            |-- painel-comercio.html\n'
    '            |-- lista-produtos.html\n'
    '            |-- cadastro-produto.html\n'
    '            |-- mercados.html\n'
    '            |-- cadastro-mercado.html\n'
    '            |-- ranking.html\n'
    '            |-- comparar.html\n'
    '            |-- lista-compra.html\n'
    '            |-- mapa.html'
)

# ── Rodapé ──────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Documentacao gerada em 31/03/2026 — MercadoCerto v1.0')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Salvar ──────────────────────────────────────────────
output_path = r'C:\Users\david.6961\mercadocerto\MercadoCerto-main\startup\DOCUMENTACAO_MercadoCerto.docx'
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
